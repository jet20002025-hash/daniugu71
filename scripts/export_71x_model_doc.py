#!/usr/bin/env python3
"""导出 71 倍模型说明为 Word 文档"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    doc = Document()
    doc.add_heading("71倍模型", 0)

    # 一、概述
    doc.add_heading("一、概述", level=1)
    doc.add_paragraph(
        "71倍模型是基于规则的技术选股模型（mode3），通过均线多头、放量突破等条件筛选强势股启动点。"
        "模型为纯规则型，无机器学习权重文件，所有评分逻辑硬编码在代码中。"
    )

    # 二、整体架构
    doc.add_heading("二、整体架构", level=1)
    doc.add_paragraph("模型分为四层：输入层 → 规则层 → 评分层 → 输出层")
    doc.add_paragraph(
        "输入层：股票池 + K线数据\n"
        "规则层：信号过滤（均线、放量等条件）\n"
        "评分层：对通过过滤的股票打分排序\n"
        "输出层：选股列表（如每日 top1/top3）",
        style="List Bullet"
    )

    # 三、输入层
    doc.add_heading("三、输入层（数据）", level=1)
    table = doc.add_table(rows=5, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "组件"
    hdr[1].text = "路径"
    hdr[2].text = "说明"
    table.rows[1].cells[0].text = "股票池"
    table.rows[1].cells[1].text = "data/gpt/stock_list.csv"
    table.rows[1].cells[2].text = "待筛选股票列表"
    table.rows[2].cells[0].text = "K线"
    table.rows[2].cells[1].text = "data/gpt/kline_cache_tencent/{code}.csv"
    table.rows[2].cells[2].text = "日线 OHLCV"
    table.rows[3].cells[0].text = "指数（可选）"
    table.rows[3].cells[1].text = "data/gpt/index_sh000001.csv"
    table.rows[3].cells[2].text = "大盘对比"
    table.rows[4].cells[0].text = "市值（可选）"
    table.rows[4].cells[1].text = "data/gpt/market_cap.csv"
    table.rows[4].cells[2].text = "市值过滤"

    # 四、规则层
    doc.add_heading("四、规则层（信号过滤）", level=1)
    doc.add_paragraph("函数：_signals_mode3 / _mode3_signals")
    doc.add_paragraph("以下条件需全部满足才产生信号：")
    table2 = doc.add_table(rows=7, cols=2)
    table2.style = "Table Grid"
    table2.rows[0].cells[0].text = "条件"
    table2.rows[0].cells[1].text = "说明"
    table2.rows[1].cells[0].text = "均线多头"
    table2.rows[1].cells[1].text = "MA10 > MA20 > MA60"
    table2.rows[2].cells[0].text = "均线斜率"
    table2.rows[2].cells[1].text = "MA10、MA20、MA60 近 3 日斜率 > 0"
    table2.rows[3].cells[0].text = "收盘价"
    table2.rows[3].cells[1].text = "收盘价 ≥ MA20"
    table2.rows[4].cells[0].text = "放量"
    table2.rows[4].cells[1].text = "成交量 ≥ 20日均量 × 1.2"
    table2.rows[5].cells[0].text = "20日涨幅"
    table2.rows[5].cells[1].text = "20日涨幅 ≤ 25%"
    table2.rows[6].cells[0].text = "历史长度"
    table2.rows[6].cells[1].text = "至少 60 根 K 线"
    # 新增行：近一年涨幅
    table2.add_row()
    table2.rows[7].cells[0].text = "近一年最高/最低"
    table2.rows[7].cells[1].text = "买点前240日内最高价/最低价 ≤ 4倍（超4倍则排除）"

    # 五、评分层
    doc.add_heading("五、评分层（打分）", level=1)
    doc.add_paragraph("函数：_score_mode3")
    doc.add_paragraph("公式：score = base(60) + ma10_ma20 + ma20_ma60 + vol_ratio + close_gap")
    doc.add_paragraph("")
    table3 = doc.add_table(rows=6, cols=3)
    table3.style = "Table Grid"
    table3.rows[0].cells[0].text = "维度"
    table3.rows[0].cells[1].text = "阈值"
    table3.rows[0].cells[2].text = "加分"
    table3.rows[1].cells[0].text = "base"
    table3.rows[1].cells[1].text = "—"
    table3.rows[1].cells[2].text = "60"
    table3.rows[2].cells[0].text = "ma10-ma20 间距"
    table3.rows[2].cells[1].text = "≥2% / ≥1% / ≥0.5%"
    table3.rows[2].cells[2].text = "+10 / +6 / +3"
    table3.rows[3].cells[0].text = "ma20-ma60 间距"
    table3.rows[3].cells[1].text = "≥2% / ≥1% / ≥0.5%"
    table3.rows[3].cells[2].text = "+10 / +6 / +3"
    table3.rows[4].cells[0].text = "放量倍数"
    table3.rows[4].cells[1].text = "≥1.6x / ≥1.4x / ≥1.2x"
    table3.rows[4].cells[2].text = "+15 / +10 / +6"
    table3.rows[5].cells[0].text = "距 MA20"
    table3.rows[5].cells[1].text = "≥3% / ≥1%"
    table3.rows[5].cells[2].text = "+5 / +3"
    doc.add_paragraph("理论分数区间：60～101（实际多为 60～100）")

    # 六、排序与输出
    doc.add_heading("六、排序与输出", level=1)
    doc.add_paragraph("同分时排序键：close_gap ↑ → vol_ratio ↓ → (ma20_gap+ma60_gap) ↓ → ret20 ↑ → code ↑")
    doc.add_paragraph("输出：按日期、按分数排序的选股列表（如每日 top1 / top3）")

    # 七、代码分布
    doc.add_heading("七、代码分布", level=1)
    table4 = doc.add_table(rows=7, cols=3)
    table4.style = "Table Grid"
    table4.rows[0].cells[0].text = "模块"
    table4.rows[0].cells[1].text = "文件"
    table4.rows[0].cells[2].text = "职责"
    table4.rows[1].cells[0].text = "信号检测"
    table4.rows[1].cells[1].text = "app/scanner.py _mode3_signals"
    table4.rows[1].cells[2].text = "规则过滤"
    table4.rows[2].cells[0].text = "信号检测"
    table4.rows[2].cells[1].text = "scripts/backtest_startup_modes.py _signals_mode3"
    table4.rows[2].cells[2].text = "回测用信号"
    table4.rows[3].cells[0].text = "评分"
    table4.rows[3].cells[1].text = "app/scanner.py _score_mode3"
    table4.rows[3].cells[2].text = "打分"
    table4.rows[4].cells[0].text = "评分"
    table4.rows[4].cells[1].text = "scripts/backtest_startup_modes.py _score_mode3"
    table4.rows[4].cells[2].text = "回测用打分"
    table4.rows[5].cells[0].text = "评分"
    table4.rows[5].cells[1].text = "app/stock_score.py score_stock"
    table4.rows[5].cells[2].text = "个股评分"
    table4.rows[6].cells[0].text = "回测"
    table4.rows[6].cells[1].text = "scripts/backtest_startup_modes.py"
    table4.rows[6].cells[2].text = "71 倍回测主脚本"

    # 八、小结
    doc.add_heading("八、小结", level=1)
    doc.add_paragraph(
        "• 类型：规则模型，无 ML 权重\n"
        "• 流程：股票池 → K线 → 规则过滤 → 打分 → 排序 → 输出\n"
        "• 权重：硬编码在代码中，无独立权重文件\n"
        "• 配置：data/models/mode3top3v1.0.json 仅作公式说明，不参与计算",
        style="List Bullet"
    )

    out_path = "data/results/71倍模型.docx"
    import os
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    print(f"已导出: {out_path}")


if __name__ == "__main__":
    main()
